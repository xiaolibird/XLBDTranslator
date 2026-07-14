# -*- coding: utf-8 -*-
"""样式化 docx 回归测试：单元格着色 / 句级三色 run / 生成有效 docx。"""
from datetime import date

from docx import Document
from docx.oxml.ns import qn

from src.scholar.schema import (
    PaperSegment, PaperMetadata, FilterDecision,
    CloseReading, CloseReadSection, CloseReadSentence,
)
from src.scholar import docx_builder


def _seg(pid, title, prio, tag=None):
    cr = CloseReading(sections=[CloseReadSection(heading="方法与数据", sentences=[
        CloseReadSentence(text="用了掩码嵌入。", tag=tag),
        CloseReadSentence(text="数据来自 MIMIC。", tag=None)])]) if tag else None
    return PaperSegment(
        segment_id=1, paper_id=pid, priority_score=prio,
        metadata=PaperMetadata(paper_id=pid, title=title, authors=["Jane Public"],
                               journal="npj DM", doi="10.1/{}".format(pid)),
        filter_decision=FilterDecision(paper_id=pid, title=title, verdict="included",
                                       decision="INCLUDE", stage="llm_judge", one_line="用处"),
        close_reading=cr)


def test_build_docx_creates_shaded_table_and_colored_runs(tmp_path):
    segs = [_seg("p1", "High Priority MNAR", 0.9, tag="方法学创新"),
            _seg("p2", "Low Priority Paper", 0.2)]
    out = docx_builder.build_digest_docx(
        segs, {"p1": "public2025a", "p2": "public2025b"},
        tmp_path / "d.docx", title="Digest",
        csl_items=[{"id": "public2025a", "title": "T", "author": [{"family": "Public", "given": "Jane"}],
                    "issued": {"date-parts": [[2025, 3, 1]]}, "DOI": "10.1/p1"}])
    assert out.exists()

    doc = Document(str(out))
    # 有优先级速览表
    assert len(doc.tables) >= 1
    table = doc.tables[0]
    # 表头 + 2 行数据
    assert len(table.rows) == 3
    # 第一行数据的首格有底纹（w:shd fill 非白）
    data_cell = table.rows[1].cells[0]
    shd = data_cell._tc.get_or_add_tcPr().find(qn("w:shd"))
    assert shd is not None and shd.get(qn("w:fill")) not in (None, "FFFFFF", "auto")


def test_docx_sentence_tag_gets_color(tmp_path):
    seg = _seg("p1", "Tagged", 0.9, tag="重要发现")
    out = docx_builder.build_digest_docx([seg], {"p1": "x2025"}, tmp_path / "d.docx")
    doc = Document(str(out))
    # 收集所有带颜色的 run，应包含紫色(800080)——重要发现
    colored = []
    for p in doc.paragraphs:
        for r in p.runs:
            if r.font.color and r.font.color.rgb is not None:
                colored.append(str(r.font.color.rgb))
    assert "800080" in colored  # 紫=重要发现


def test_docx_sorts_by_priority(tmp_path):
    segs = [_seg("lo", "LOWONE", 0.1), _seg("hi", "HIGHONE", 0.9)]
    out = docx_builder.build_digest_docx(segs, {"lo": "a", "hi": "b"}, tmp_path / "d.docx")
    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    # 速览表/正文里 HIGHONE 应排在 LOWONE 之前
    assert text.index("HIGHONE") < text.index("LOWONE")
