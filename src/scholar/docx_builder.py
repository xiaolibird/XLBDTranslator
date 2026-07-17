# -*- coding: utf-8 -*-
"""
样式化 docx 生成（python-docx）：优先级速览表单元格着色、句级三色联想着色、字体区分、加宽标题列。

比 pandoc→docx 更可控——单元格底纹(w:shd) 和句级 run 着色 pandoc 做不到。
markdown 版仍由 notes.py 产出（供 pandoc→LaTeX）；这里是并行的样式化 docx 分支。
三色：方法学创新=墨绿 / 重要发现=紫 / 研究背景=蓝（关联研究主线）。
"""
from pathlib import Path
from typing import List, Dict, Any, Optional

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .schema import PaperSegment
from .notes import _priority_tier
from ..utils.logger import get_logger

logger = get_logger(__name__)

# tier emoji → 单元格底纹色
_TIER_FILL = {"🔴": "FFC7CE", "🟠": "FFEB9C", "🟢": "C6EFCE"}
# 句级角色 tag → 字体颜色（墨绿/紫/蓝）。新三类 + 旧三类（历史 bundle 兼容）。
_TAG_COLOR = {
    "方法论借鉴": RGBColor(0x00, 0x64, 0x00),   # 墨绿
    "可引用证据": RGBColor(0x80, 0x00, 0x80),   # 紫
    "可反驳观点": RGBColor(0x00, 0x00, 0xCD),   # 蓝
    # legacy：
    "方法学创新": RGBColor(0x00, 0x64, 0x00),
    "重要发现": RGBColor(0x80, 0x00, 0x80),
    "研究背景": RGBColor(0x00, 0x00, 0xCD),
}
_META_GRAY = RGBColor(0x66, 0x66, 0x66)


def _shade_cell(cell, hex_fill: str):
    """给表格单元格加底纹（注入 w:shd）。"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _apply_font(run, size: Optional[float] = None, bold: Optional[bool] = None,
                color: Optional[RGBColor] = None, cjk: str = ""):
    """设置 run 字体（含 CJK eastAsia），字号/加粗/颜色。"""
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if cjk:
        run.font.name = cjk
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:eastAsia"), cjk)


def _set_col_widths(table, widths_inches: List[float]):
    """固定表格列宽（关 autofit，逐格设宽）。"""
    from docx.shared import Inches
    table.autofit = False
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_inches):
                cell.width = Inches(widths_inches[i])


def build_digest_docx(segments: List[PaperSegment], citekeys: Dict[str, Optional[str]],
                      out_path: Path, title: str = "Scholar Digest",
                      csl_items: Optional[List[Dict[str, Any]]] = None,
                      cjk_font: str = "", instruction: str = "") -> Path:
    """生成样式化 digest docx。按优先级降序；速览表单元格着色；精读句级三色着色。"""
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    ordered = sorted(segments, key=lambda s: s.priority_score, reverse=True)
    total = len(ordered)
    tiers = [_priority_tier(i, total) for i in range(total)]

    doc = Document()

    # 标题
    h = doc.add_heading(title, level=0)
    for r in h.runs:
        _apply_font(r, cjk=cjk_font)
    p = doc.add_paragraph()
    r = p.add_run("共 {} 篇（按优先级降序）".format(total))
    _apply_font(r, size=10.5, color=_META_GRAY, cjk=cjk_font)

    # 三色图例
    lp = doc.add_paragraph()
    _apply_font(lp.add_run("句级联想标记："), size=9, color=_META_GRAY, cjk=cjk_font)
    for name, col in _TAG_COLOR.items():
        _apply_font(lp.add_run(" ■{} ".format(name)), size=9, color=col, cjk=cjk_font)

    # ---- 优先级速览表 ----
    doc.add_heading("优先级速览", level=1)
    cols = ["优先级", "#", "裁决", "标题", "一句话用处"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Light Grid Accent 1"
    for i, c in enumerate(cols):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(c)
        _apply_font(run, size=10, bold=True, cjk=cjk_font)
    for i, seg in enumerate(ordered):
        fd = seg.filter_decision
        row = table.add_row().cells
        emoji = tiers[i].split()[0]
        vals = [tiers[i], str(i + 1),
                (fd.decision if fd and fd.decision else ""),
                (seg.metadata.title or ""),
                (fd.one_line if fd and fd.one_line else "")]
        for j, v in enumerate(vals):
            run = row[j].paragraphs[0].add_run(v)
            _apply_font(run, size=9.5, cjk=cjk_font)
        _shade_cell(row[0], _TIER_FILL.get(emoji, "FFFFFF"))
    _set_col_widths(table, [0.7, 0.35, 0.6, 3.6, 1.7])  # 标题列加宽

    # ---- 逐篇 ----
    for i, seg in enumerate(ordered):
        meta = seg.metadata
        ck = citekeys.get(seg.paper_id)
        head = doc.add_heading(level=2)
        _apply_font(head.add_run("{} {}. {}".format(tiers[i], i + 1, meta.title or "(无标题)")),
                    cjk=cjk_font)

        # 元信息（小字灰）
        mp = doc.add_paragraph()
        bits = []
        if meta.authors:
            bits.append("作者: {}{}".format(", ".join(meta.authors[:5]),
                                          " et al." if len(meta.authors) > 5 else ""))
        if meta.journal:
            bits.append("期刊: {}".format(meta.journal))
        if meta.doi:
            bits.append("DOI: {}".format(meta.doi))
        if ck:
            bits.append("citekey: {}".format(ck))
        if seg.priority_score:
            bits.append("优先级: {:.2f}".format(seg.priority_score))
        _apply_font(mp.add_run("  ·  ".join(bits)), size=9, color=_META_GRAY, cjk=cjk_font)

        cr = seg.close_reading
        if cr and cr.sections:
            # 全文精读（句级三色）
            label = "全文精读" if cr.from_full_text else "精读（仅摘要）"
            _apply_font(doc.add_paragraph().add_run(label), size=10.5, bold=True, cjk=cjk_font)
            for sec in cr.sections:
                sp = doc.add_paragraph()
                _apply_font(sp.add_run("【{}】".format(sec.heading)), size=10, bold=True, cjk=cjk_font)
                bp = doc.add_paragraph()
                for st in sec.sentences:
                    run = bp.add_run(st.text + " ")
                    _apply_font(run, size=10.5, color=_TAG_COLOR.get(st.tag), cjk=cjk_font)
        else:
            # 无精读 → 摘要 + AI 总结
            _apply_font(doc.add_paragraph().add_run("摘要"), size=10.5, bold=True, cjk=cjk_font)
            _apply_font(doc.add_paragraph().add_run(
                seg.translated_abstract or seg.original_abstract or "（摘要暂无）"),
                size=10.5, cjk=cjk_font)
            if seg.summary:
                _apply_font(doc.add_paragraph().add_run("AI 归纳"), size=10.5, bold=True, cjk=cjk_font)
                _apply_font(doc.add_paragraph().add_run(seg.summary), size=10.5, cjk=cjk_font)

    # ---- 参考文献 ----
    if csl_items:
        doc.add_heading("参考文献", level=1)
        for it in csl_items:
            authors = "; ".join(
                "{} {}".format(a.get("family", ""), a.get("given", "")).strip()
                for a in (it.get("author") or [])[:6])
            year = ""
            dp = (it.get("issued") or {}).get("date-parts") or []
            if dp and dp[0]:
                year = str(dp[0][0])
            parts = [x for x in [authors, "({})".format(year) if year else "",
                                 it.get("title", ""), it.get("container-title", ""),
                                 "DOI: {}".format(it["DOI"]) if it.get("DOI") else ""] if x]
            _apply_font(doc.add_paragraph().add_run(". ".join(parts)), size=9.5, cjk=cjk_font)

    doc.save(str(out_path))
    logger.info("  📄 样式化 docx → {}".format(out_path))
    return out_path
